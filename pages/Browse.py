import streamlit as st
import pandas as pd 
import json
from io import BytesIO
from docx import Document
from database.model import get_all_id_and_filename, get_transcription, get_feature, get_report, regenerate_report_col
from Utils.utils import make_text_readable_in_markdown

'''
Streamlit page for Browse transcriptions

This page allows users to select a transcription from a dropdown menu and view the transcription and download transcript and report data.
'''

def display_keywords(words):
    for i in range(0, len(words), 5):
        cols = st.columns(5)  
        for col, word in zip(cols, words[i:i+5]):
            with col:
                st.markdown(f"""
                    <div style="background-color: #f9f9f9;
                                border-radius: 10px;
                                padding: 10px;
                                box-shadow: 0px 2px 10px rgba(0, 0, 0, 0.1);
                                text-align: center;">
                        <span style="font-size: 18px; color: #333;">{word}</span>
                    </div>
                """, unsafe_allow_html=True)

# Function to generate transcript docx file
def generate_transcript_docx(transcription_text, transcript_id):
    doc = Document()
    doc.add_heading(f'Transcription ID: {transcript_id}', 0)
    doc.add_paragraph(transcription_text)
    return doc

# Function to generate report docx file
def generate_report_docx(summary_data, background_data, reflection_data, conclusion_data, keywords_data, transcript_id):
    doc = Document()
    doc.add_heading(f'Report for Transcription ID: {transcript_id}', 0)
    
    doc.add_heading('Background', level=1)
    doc.add_paragraph(background_data[0])
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(summary_data[0])
    
    doc.add_heading('Reflection', level=1)
    doc.add_paragraph(reflection_data[0])
    
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(conclusion_data[0])
    
    doc.add_heading('Keywords', level=1)
    if isinstance(keywords_data, list):
        doc.add_paragraph(", ".join(keywords_data))  # Joining keywords as a string
    else:
        doc.add_paragraph(str(keywords_data))  # Handling other cases
    
    return doc

# Convert DOCX to downloadable bytes
def convert_to_downloadable(doc):
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

# Function to display keywords as cards
def display_keywords(words):
    # Create a number of columns equal to the number of words
    cols = st.columns(len(words))
    
    # Loop over each word and display it in a column
    for col, word in zip(cols, words):
        with col:
            st.markdown(f"""
                <div style="background-color: #ADD8E6;
                            border-radius: 10px;
                            padding: 10px;
                            box-shadow: 0px 2px 10px rgba(0, 0, 0, 0.1);
                            text-align: center;">
                    <span style="font-size: 18px; color: #333;">{word}</span>
                </div>
            """, unsafe_allow_html=True)
# Setting page config
st.set_page_config(page_title="Transcription Application", page_icon=":microphone:", layout="wide", initial_sidebar_state="expanded")
st.title("Browse transcriptions")

# Fetch all transcription data
ALL_TRANSCRIPTIONS_TUPLE = get_all_id_and_filename()
ALL_TRANSCRIPTIONS_LIST = [f"{pair[0]}: {pair[1]}" for pair in ALL_TRANSCRIPTIONS_TUPLE]

# Dropdown for selecting a transcription
option = st.selectbox("Select a transcription to view", ALL_TRANSCRIPTIONS_LIST)

if option:
    st.markdown(f"### Selected transcription: {option}")

    id = int(option.split(":")[0])
    
    # Fetch the transcription and make it readable
    transcription = get_transcription(id)
    readable_transcription = make_text_readable_in_markdown(transcription[0][0], id)
    

    # Fetch features and report data
    keywords = get_feature(id, "Keywords")
    summary = get_report(id, "Summary")
    background = get_report(id, "Background")
    reflection = get_report(id, "Reflection")
    conclusion = get_report(id, "Conclusion")
    
    # Data formatting for reports
    summary_data = [summary[0][0]] if summary else [""]
    background_data = [background[0][0]] if background else [""]
    reflection_data = [reflection[0][0]] if reflection else [""]
    conclusion_data = [conclusion[0][0]] if conclusion else [""]
    keywords_data = [keywords[0][0]] if keywords else [""]

    col1, col2 = st.columns(2)
    
    with col1:
        # Download button for transcript
        transcript_doc = generate_transcript_docx(readable_transcription, id)
        transcript_file = convert_to_downloadable(transcript_doc)
        st.download_button(
            label="Download Transcript",
            data=transcript_file,
            file_name=f"transcription_{id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    with col2:
        # Download button for report
        report_doc = generate_report_docx(summary_data, background_data, reflection_data, conclusion_data, keywords_data, id)
        report_file = convert_to_downloadable(report_doc)
        st.download_button(
            label="Download Report",
            data=report_file,
            file_name=f"report_{id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    st.markdown(f"##### Transcript : {option}")
    st.markdown(f"""
    <div style="
        background-color: #ADD8E6;
        border: 1px solid #000000;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
    ">
        <p style="color: #333;">{readable_transcription}</p>
    </div>
    """, unsafe_allow_html=True)
    # Display report table
    st.markdown(f"##### Report : {option}")
    report_df = pd.DataFrame({
        "Background": background_data,
        "Summary": summary_data,
        "Reflection": reflection_data,
        "Conclusion": conclusion_data,
    })
    st.dataframe(report_df)


    st.markdown(f"##### Keywords")
    try:
        if keywords_data and isinstance(keywords_data[0], str):
            keyword_json = json.loads(keywords_data[0].replace("'", "\""))  # Handle single quotes for JSON parsing
        else:
            keyword_json = {}
        title = "Keywords" if "Keywords" in keyword_json else "keywords"
        display_keywords(keyword_json.get(title, []))
    except Exception as e:
        st.warning("Failed to load keywords.")
        st.dataframe(pd.DataFrame({"Keywords": keywords_data}))

    # Columns for download buttons
    
